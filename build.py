import csv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

document = Document()

#set styles 


with open('index1.tsv','r') as original_index:
    index_ptr = csv.reader(original_index, delimiter='\t')
    #sort alphabetically by first word 
    data = (list(index_ptr))
    tracker=''

    for line in data:
        current = line[0][:1]
        if current.isalpha() and tracker.upper() != current.upper():
            letter_header = document.add_paragraph().add_run(line[0][:1])
            letter_font=letter_header.font
            letter_font.color.rgb = RGBColor(255, 165, 0)
            letter_font.name='Calibri(Headings)'
            letter_font.bold=True
            letter_font.size=Pt(20)

        para = document.add_paragraph('')
        #format term
        para.add_run(line[0], style="Heading 2 Char")
        #format book.page number
        para.add_run(" [")
        para.add_run(line[1]).italic = True
        para.add_run("] ")
        para.add_run("\n")
        #format definiton/context
        para.add_run(line[2])

        tracker = line[0][:1]

#format with two columns
section = document.sections[0]
section_ptr = section._sectPr
columns = section_ptr.xpath('./w:cols')[0]
columns.set(qn('w:num'),'2')


document.save('index_finalized.docx')
